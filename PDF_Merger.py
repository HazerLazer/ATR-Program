from __future__ import annotations

from pathlib import Path
from typing import Dict, List, Tuple, Union
import re

from pypdf import PdfReader, PdfWriter   # pip install pypdf

from docx import Document
import os
import sys
import win32com.client as win32
import re
import shutil

from concurrent.futures import ThreadPoolExecutor, as_completed

# ---------------------------------------------------------------------
# Regex that grabs the leading digits of a section label:
#   "3.2A"  →  "3"
#   "5Footnotes"  →  "5"
SEC_NUM_RE = re.compile(r"^(\d+)")
PARSE_RE = re.compile(
    r"""^T_           # literal "T_"
        (?P<num>[^_]+)  # tif-num is everything up to the next underscore
        _               # underscore separator
        (?P<name>.+?)   # tif-name: one or more chars, reluctantly
        AR
        (?P<year>\d+)   # year_suffix: one or more digits
        -
        (?P<section>.+) # section: the rest
        $""",
    re.VERBOSE
)

# ---------------------------------------------------------------------
def _parse(stem: str) -> Tuple[str, str, str] | None:
    """
    Split a filename stem
        T_<num>_<name>_<section>
    into (tif_num, tif_name, section).

    <name> can itself contain underscores, so everything between the
    2nd and final underscore is treated as the name.

    Returns None if the stem doesn’t match the pattern.
    """
    m = PARSE_RE.match(stem)
    if not m:
        return None

    tif_num   = m.group("num")
    tif_name  = m.group("name")          # does *not* include AR<year>
    # year_suf  = m.group("year")          # only if you need to use it
    section   = m.group("section")       # e.g. "3.2A" or "7"

    # if you want the year baked into the name for grouping:
    # tif_name = f"{tif_name}AR{year_suf}"

    return tif_num, tif_name, section

def merge_by_tif_number(
    directory: Union[str, Path],
    year_suffix: str,
    recursive: bool = False,
    out_suffix: str = "-1.pdf",
    delete_sources: bool = False,
    all: bool = False,
    max_workers: int | None = None,
) -> List[Path]:
    """
    Merge section PDFs for each (tif-num, tif-name) in parallel using threads.
    """
    directory = Path(directory).expanduser().resolve()
    pattern = "**/*.pdf" if recursive else "*.pdf"

    # 1) Bucket all PDFs by (tif_num, tif_name)
    buckets: Dict[Tuple[str, str], List[Path]] = {}
    for pdf in directory.glob(pattern):
        if not pdf.is_file():
            continue
        parsed = _parse(pdf.stem)
        if not parsed:
            continue
        tif_num, tif_name, section = parsed

        if not all:
            m = SEC_NUM_RE.match(section)
            if not m:
                continue
            num = int(m.group(1)); rest = section[m.end():]
            if num >= 8 or (num == 7 and rest):
                continue

        buckets.setdefault((tif_num, tif_name), []).append(pdf)

    if not buckets:
        raise FileNotFoundError(f"No qualifying PDFs found under {directory}")

    written: List[Path] = []

    # 2) Define a per-bucket merge function
    def _merge_one(tif_num: str, tif_name: str, pdf_list: List[Path]) -> Path:
        pdf_list.sort(key=lambda p: _parse(p.stem)[2])
        writer = PdfWriter()
        for pdf in pdf_list:
            reader = PdfReader(pdf)
            for page in reader.pages:
                writer.add_page(page)

        out_path = directory / f"T_{tif_num}_{tif_name}AR{year_suffix}{out_suffix}"
        with out_path.open("wb") as fh:
            writer.write(fh)

        if delete_sources:
            for src in pdf_list:
                try:
                    if not src.samefile(out_path):
                        src.unlink()
                except Exception:
                    pass

        return out_path

    # 3) Merge in parallel using a thread pool
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_key = {
            executor.submit(_merge_one, tif_num, tif_name, lst): (tif_num, tif_name)
            for (tif_num, tif_name), lst in buckets.items()
        }
        for fut in as_completed(future_to_key):
            key = future_to_key[fut]
            try:
                result_path = fut.result()
                print(f"created {result_path.relative_to(directory)} for {key}")
                written.append(result_path)
            except Exception as exc:
                print(f"Error merging {key}: {exc}")

    return written


# def _merge_one(
#     tif_num: str,
#     tif_name: str,
#     pdf_list: List[Path],
#     directory: Path,
#     year_suffix: str,
#     out_suffix: str,
#     delete_sources: bool,
# ) -> Path:
#     """Worker for merging a single bucket."""
#     # sort pages by section
#     pdf_list.sort(key=lambda p: _parse(p.stem)[2])

#     writer = PdfWriter()
#     for pdf in pdf_list:
#         reader = PdfReader(pdf)
#         for page in reader.pages:
#             writer.add_page(page)

#     out_path = directory / f"T_{tif_num}_{tif_name}AR{year_suffix}{out_suffix}"
#     with out_path.open("wb") as fh:
#         writer.write(fh)

#     if delete_sources:
#         for pdf in pdf_list:
#             try:
#                 if not pdf.samefile(out_path):
#                     pdf.unlink()
#             except Exception:
#                 pass

#     return out_path

# def merge_by_tif_number(
#     directory: Union[str, Path],
#     year_suffix=None,
#     recursive: bool = False,
#     out_suffix: str = "-1.pdf",
#     delete_sources: bool = False,
#     all: bool = False,
#     max_workers: int = 2,     # new param to control parallelism
# ) -> List[Path]:
#     directory = Path(directory).expanduser().resolve()
#     pattern = "**/*.pdf" if recursive else "*.pdf"
#     buckets: Dict[Tuple[str, str], List[Path]] = {}

#     # 1) Bucket the PDFs (serial I/O)
#     for pdf in directory.glob(pattern):
#         if not pdf.is_file():
#             continue
#         parsed = _parse(pdf.stem)
#         if not parsed:
#             continue
#         tif_num, tif_name, section = parsed
#         if not all:
#             m = SEC_NUM_RE.match(section)
#             if not m:
#                 continue
#             num = int(m.group(1))
#             rest = section[m.end():]
#             if num >= 8 or (num == 7 and rest):
#                 continue
#         buckets.setdefault((tif_num, tif_name), []).append(pdf)

#     if not buckets:
#         raise FileNotFoundError(f"No qualifying PDFs found under {directory}")

#     written: List[Path] = []

#     # 2) Merge in parallel
#     with ProcessPoolExecutor(max_workers=max_workers) as exe:
#         futures = {
#             exe.submit(
#                 _merge_one,
#                 tif_num,
#                 tif_name,
#                 pdf_list,
#                 directory,
#                 year_suffix,
#                 out_suffix,
#                 delete_sources,
#             ): (tif_num, tif_name)
#             for (tif_num, tif_name), pdf_list in buckets.items()
#         }
#         for fut in as_completed(futures):
#             tif_key = futures[fut]
#             try:
#                 out_path = fut.result()
#                 print(f"created {out_path.relative_to(directory)} for {tif_key}")
#                 written.append(out_path)
#             except Exception as e:
#                 print(f"Error merging {tif_key}: {e}")

#     return written





# def merge_by_tif_number(
#     directory: Union[str, Path],
#     year_suffix = None,
#     recursive: bool = False,
#     out_suffix: str = "-1.pdf",
#     delete_sources: bool = False,
#     all: bool = False,
# ) -> List[Path]:
#     """
#     Merge section PDFs for each (tif-num, tif-name) found in *directory*.

#     Parameters
#     ----------
#     directory : str | Path
#         Folder that holds the PDFs.
#     recursive : bool, default False
#         If True, also scan sub-folders.
#     out_suffix : str, default '_1.pdf'
#         Suffix for the merged file (client wants '_1.pdf').
#     delete_sources : bool, default False
#         If True, delete the individual section PDFs after merging.
#         The merged PDF itself is never deleted.
#     all : bool, default False
#         If True, don't skip any pdfs

#     Returns
#     -------
#     list[Path]
#         Paths of merged PDFs that were written.
#     """
#     directory = Path(directory).expanduser().resolve()
#     pattern = "**/*.pdf" if recursive else "*.pdf"

#     buckets: Dict[Tuple[str, str], List[Path]] = {}

#     for pdf in directory.glob(pattern):
#         if not pdf.is_file():
#             continue
#         parsed = _parse(pdf.stem)
#         if not parsed:
#             continue
#         tif_num, tif_name, section = parsed

#         if not all:
#             m = SEC_NUM_RE.match(section)
#             if not m:
#                 continue       # no numeric prefix – skip
#             num = int(m.group(1))
#             rest = section[m.end():] 
#             if num >= 8 or (num == 7 and rest):
#                 continue       # drop Section 7.1, 8, 9, …

#         buckets.setdefault((tif_num, tif_name), []).append(pdf)
        
#     if not buckets:
#         raise FileNotFoundError(f"No qualifying PDFs found under {directory}")

#     written: List[Path] = []

#     for (tif_num, tif_name), pdf_list in buckets.items():
#         pdf_list.sort(key=lambda p: _parse(p.stem)[2])

#         writer = PdfWriter()
#         for pdf in pdf_list:
#             reader = PdfReader(pdf)
#             for page in reader.pages:
#                 writer.add_page(page)

#         out_path = directory / f"T_{tif_num}_{tif_name}AR{year_suffix}{out_suffix}"
#         with out_path.open("wb") as fh:
#             writer.write(fh)
#         print(f"created {out_path.relative_to(directory)}")
#         written.append(out_path)

#         if delete_sources:
#             for pdf in pdf_list:
#                 try:
#                     if pdf.samefile(out_path):   # don’t delete merged PDF
#                         continue
#                     pdf.unlink()
#                 except Exception as exc:
#                     print(f"couldn't delete {pdf.name}: {exc}")

#     return written


def bc_docs(output_dir, name, tif_name, attB_ijrl_file, attC_ijrl_file, attB_tifcorp_file, 
            attC_tifcorp_file, bsigned_file, csigned_file, ijrl):
    word_app = win32.Dispatch("Word.Application")
    word_app.Visible = False
    word_app.DisplayAlerts = 0
    
    wdExportFormatPDF           = 17
    wdExportOptimizeForPrint    = 0
    wdExportFromTo              = 3
    wdExportDocumentContent     = 0
    wdExportCreateWordBookmarks = 0
    
    if ijrl:
        att_B = Document(attB_ijrl_file)
        att_C = Document(attC_ijrl_file)
    else:
        att_B = Document(attB_tifcorp_file)
        att_C = Document(attC_tifcorp_file)
        
    for paragraph in att_B.paragraphs:
        if "{{name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{name}}", name)

    for paragraph in att_C.paragraphs:
        if "{{name}}" in paragraph.text:
            paragraph.text = paragraph.text.replace("{{name}}", name)
            
    name1 = tif_name + "-" + "B"
    safe_name1 = re.sub(r'[\\/:\*\?"<>\|]', '_', name1)
    temp_word_path = os.path.join(output_dir, f"Temp_{safe_name1}.docx")
    temp_word_path = os.path.normpath(temp_word_path)

    att_B.save(temp_word_path)

    name2 = tif_name + "-" + "C"
    safe_name2 = re.sub(r'[\\/:\*\?"<>\|]', '_', name2)
    temp_word_path2 = os.path.join(output_dir, f"Temp_{safe_name2}.docx")
    temp_word_path2 = os.path.normpath(temp_word_path2)

    att_C.save(temp_word_path2)
    
    pdf_output_path = os.path.join(output_dir, f"{safe_name1}.pdf")
    pdf_output_path = os.path.normpath(pdf_output_path)

    pdf_output_path2 = os.path.join(output_dir, f"{safe_name2}.pdf")
    pdf_output_path2 = os.path.normpath(pdf_output_path2)

    from win32com.client import constants

    # Open Temp_<safe_name1>.docx
    word_doc = word_app.Documents.Open(temp_word_path)

    # Export page 1 only:
    #   - ExportFormat = constants.wdExportFormatPDF  (17)
    #   - ExportRange  = constants.wdExportFromTo    (3)
    #   - From = 1, To = 1
    #   - OpenAfterExport=False, OptimizeFor=constants.wdExportOptimizeForPrint (0), etc.
    word_doc.ExportAsFixedFormat(
        OutputFileName=pdf_output_path,
        ExportFormat=wdExportFormatPDF,
        OpenAfterExport=False,
        OptimizeFor=wdExportOptimizeForPrint,
        Range=wdExportFromTo,
        From=1,
        To=1,
        Item=wdExportDocumentContent,
        CreateBookmarks=wdExportCreateWordBookmarks,
        DocStructureTags=True,
        BitmapMissingFonts=True,
        UseISO19005_1=False
    )
    word_doc.Close(False)  # close without saving changes

    # Repeat for Temp_<safe_name2>.docx → only page 1
    word_doc2 = word_app.Documents.Open(temp_word_path2)
    word_doc2.ExportAsFixedFormat(
        OutputFileName=pdf_output_path2,
        ExportFormat=wdExportFormatPDF,
        OpenAfterExport=False,
        OptimizeFor=wdExportOptimizeForPrint,
        Range=wdExportFromTo,
        From=1,
        To=1,
        Item=wdExportDocumentContent,
        CreateBookmarks=wdExportCreateWordBookmarks,
        DocStructureTags=True,
        BitmapMissingFonts=True,
        UseISO19005_1=False
    )
    word_doc2.Close(False)

    os.remove(temp_word_path)
    os.remove(temp_word_path2)
    
    dst_b = os.path.join(output_dir, f"{safe_name1}1.pdf")
    try:
        shutil.copy(bsigned_file, dst_b)
    except:
        print("didn't print B")
        return
    
    dst_c = os.path.join(output_dir, f"{safe_name2}1.pdf")
    try:
        shutil.copy(csigned_file, dst_c)
    except:
        print("didn't print C")
        return
    
    word_app.Quit()
    
    return
    